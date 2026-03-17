"""
DOCX to TEITOK-style TEI conversion.

Produces TEI with header, paragraphs, <hi> styles,
tables, images, hyperlinks, footnotes). Used by load_docx; the resulting TEI root
can be stored in the Document and written verbatim by save_teitok for identical output.
"""
from __future__ import annotations

import copy
import os
import tempfile
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from lxml import etree

from ..core.model import Document

import re

# Wordprocessing/drawing namespaces for DOCX
NAMESPACES = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


def _require_docx():
    try:
        from docx import Document as DocxDocument  # type: ignore
        from docx.opc.exceptions import PackageNotFoundError  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires the 'python-docx' package. "
            "Install with: pip install 'flexiconv[docx]'"
        ) from exc
    return DocxDocument, PackageNotFoundError


def _qn(prefix: str, local: str) -> str:
    return "{" + NAMESPACES[prefix] + "}" + local


def _get_tag(element: etree._Element) -> str:
    tag = element.tag
    for ns in NAMESPACES:
        uri = NAMESPACES[ns]
        tag = tag.replace("{" + uri + "}", "")
    return tag


def _extract_images_and_map(
    docx_path: str,
    image_dir: str,
) -> dict[str, str]:
    """Extract images from DOCX to image_dir; return relationship_id -> filename map."""
    result: dict[str, str] = {}
    with open(docx_path, "rb") as f:
        import zipfile
        with zipfile.ZipFile(f, "r") as z:
            media_files = [n for n in z.namelist() if n.startswith("word/media/")]
            os.makedirs(image_dir, exist_ok=True)
            for name in media_files:
                filename = os.path.basename(name)
                out_path = os.path.join(image_dir, filename)
                with z.open(name) as inf:
                    with open(out_path, "wb") as outf:
                        outf.write(inf.read())
                rel_id = os.path.splitext(filename)[0]
                result[rel_id] = filename
            rels_path = "word/_rels/document.xml.rels"
            if rels_path in z.namelist():
                with z.open(rels_path) as rels_file:
                    rels_root = etree.fromstring(rels_file.read())
                for rel in rels_root.findall(
                    ".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
                ):
                    if "image" in (rel.get("Type") or ""):
                        rid = rel.get("Id")
                        target = rel.get("Target") or ""
                        if target.startswith("media/"):
                            result[rid or ""] = os.path.basename(target)
    return result


def _get_color(run: Any) -> str:
    if run.font.color and run.font.color.rgb:
        return f"color: #{run.font.color.rgb};"
    if run.style and run.style.font.color and run.style.font.color.rgb:
        return f"color: #{run.style.font.color.rgb};"
    return ""


def _get_font_size(run: Any) -> str:
    if run.font.size:
        return f"font-size: {run.font.size.pt}pt;"
    if run.style and run.style.font.size:
        return f"font-size: {run.style.font.size.pt}pt;"
    return ""


def _get_text_styles(run: Any) -> str:
    parts = []
    if run.bold or (run.style and run.style.font.bold):
        parts.append("font-weight: bold;")
    if run.italic or (run.style and run.style.font.italic):
        parts.append("font-style: italic;")
    if run.underline or (run.style and run.style.font.underline):
        parts.append("text-decoration: underline;")
    if run.font.superscript or (run.style and run.style.font.superscript):
        parts.append("vertical-align: super; font-size: smaller;")
    if run.font.subscript or (run.style and run.style.font.subscript):
        parts.append("vertical-align: sub; font-size: smaller;")
    return " ".join(parts)


def _normalize_line_height(line_spacing: Any, line_spacing_rule: Any) -> float:
    if line_spacing_rule == 4:  # WD_LINE_SPACING.EXACTLY
        return line_spacing / 20
    if line_spacing_rule == 5:  # WD_LINE_SPACING.MULTIPLE
        return line_spacing
    return 1.2


def _paragraph_to_css(para: Any) -> str:
    css = []
    align_map = {None: "left", 0: "left", 1: "center", 2: "right", 3: "justify"}
    alignment = align_map.get(para.paragraph_format.alignment, "left")
    css.append(f"text-align: {alignment};")
    if para.paragraph_format.left_indent:
        css.append(f"margin-left: {para.paragraph_format.left_indent.pt}pt;")
    if para.paragraph_format.right_indent:
        css.append(f"margin-right: {para.paragraph_format.right_indent.pt}pt;")
    if para.paragraph_format.first_line_indent:
        css.append(f"text-indent: {para.paragraph_format.first_line_indent.pt}pt;")
    if para.paragraph_format.space_before:
        css.append(f"margin-top: {para.paragraph_format.space_before.pt}pt;")
    if para.paragraph_format.space_after:
        css.append(f"margin-bottom: {para.paragraph_format.space_after.pt}pt;")
    if para.paragraph_format.line_spacing:
        lh = _normalize_line_height(
            para.paragraph_format.line_spacing,
            para.paragraph_format.line_spacing_rule,
        )
        css.append(f"line-height: {lh};")
    shd = para._element.find(".//w:shd", namespaces=NAMESPACES)
    if shd is not None:
        fill = shd.get(_qn("w", "fill"))
        if fill and fill != "FFFFFF":
            css.append(f"background-color: #{fill}")
    return " ".join(css)


def _infer_heading_level(head_el: etree._Element) -> int:
    """
    Infer a heading level (1–6) for a TEI <head> element using several cues:

    1) @type on <head>, e.g. type="main"/"sub"/"subsub" or "h1"/"h2"/"h3".
    2) @type on the enclosing <div>, e.g. type="chapter"/"section"/"subsection".
    3) Nesting depth of <div> ancestors under <body>/<text>.
    4) @n patterns like "1", "1.1", "1.1.1" → level 1, 2, 3.
    """

    def _level_from_type(val: Optional[str]) -> Optional[int]:
        if not val:
            return None
        v = val.strip().lower()
        if not v:
            return None
        # Direct h1..h6 style
        m = re.match(r"h([1-6])$", v)
        if m:
            return int(m.group(1))
        # Common semantic labels
        if v in {"main", "chapter", "title"}:
            return 1
        if v in {"sub", "section", "sec"}:
            return 2
        if v in {"subsub", "subsection", "sub-section"}:
            return 3
        return None

    # 1) @type on <head>
    lvl = _level_from_type(head_el.get("type"))
    if lvl:
        return min(max(lvl, 1), 6)

    # 2) @type on enclosing <div>
    parent = head_el.getparent()
    if parent is not None and (parent.tag or "").split("}")[-1] == "div":
        lvl = _level_from_type(parent.get("type"))
        if lvl:
            return min(max(lvl, 1), 6)

    # 3) Nesting depth of <div> under <body>/<text>
    depth = 0
    anc = parent
    while anc is not None:
        tag = (anc.tag or "").split("}")[-1]
        if tag == "div":
            depth += 1
        if tag in {"body", "text"}:
            break
        anc = anc.getparent()
    if depth > 0:
        return min(depth, 6)

    # 4) @n like "1", "1.1", "1.1.1"
    n_val = (head_el.get("n") or "").strip()
    if n_val:
        # Split on dot or other common separators.
        parts = re.split(r"[.\-]", n_val)
        nums = [p for p in parts if p.isdigit()]
        if nums:
            return min(len(nums), 6)

    # Fallback
    return 1


def _process_run(run: Any, footnote_map: dict[str, str]) -> etree._Element:
    elem = etree.Element("hi")
    styles = _get_color(run) + _get_font_size(run) + _get_text_styles(run)
    if styles:
        elem.set("style", styles.strip())
    xml_str = run._element.xml if hasattr(run._element, "xml") else etree.tostring(run._element, encoding="unicode")
    if "w:footnoteReference" in xml_str:
        run_xml = etree.fromstring(run._element.xml) if hasattr(run._element, "xml") else run._element
        for fn_ref in run_xml.findall(".//w:footnoteReference", namespaces=NAMESPACES):
            fn_id = fn_ref.get(_qn("w", "id"))
            if fn_id in footnote_map:
                elem = etree.Element("note")
                elem.set("id", "fn-" + str(fn_id))
                elem.text = footnote_map[fn_id]
    else:
        text = run.text or ""
        if "\n" not in text:
            elem.text = text
        else:
            parts = text.split("\n")
            elem.text = parts[0]
            # Each additional segment becomes an explicit line break + following text
            for extra in parts[1:]:
                lb = etree.SubElement(elem, "lb")
                if extra:
                    lb.tail = extra
    return elem


def _ensure_space_between(
    prev_tail: Optional[str],
    next_text: Optional[str],
    *,
    after_element: bool = False,
) -> str:
    """Return the string prev_tail should become so there is whitespace before next_text (for TEITOK).
    Add space when the two parts do not display as one unit. Do not add between un + real (same word).
    When after_element is True, prev_tail is empty means we are after an element (e.g. after \"1\"), so add space before \"Rus\"."""
    pt = prev_tail or ""
    nt = next_text or ""
    if not nt.strip():
        return pt
    if pt.rstrip() != pt:
        return pt  # already ends with whitespace
    if nt and nt[0].isspace():
        return pt  # next starts with whitespace
    # No previous content
    if not pt.strip():
        if after_element:
            return pt + " "  # after an element, add space before next (e.g. \"1\" then \"Rus\")
        return (pt + " ") if (nt.strip() and not nt[0].isalpha()) else pt  # start: no space before \"un\"
    # Do not add space when both sides are alphabetic (same word, e.g. un + real in un*real*)
    if pt and nt and pt[-1].isalpha() and nt[0].isalpha():
        return pt
    return pt + " "


def _ensure_trailing_space_before(parent: etree._Element, new_child: etree._Element) -> None:
    """Ensure there is whitespace between the last child of parent and new_child's text (for TEITOK)."""
    if len(parent) == 0:
        return
    last = parent[-1]
    prev_tail = last.tail or ""
    next_text = new_child.text or ""
    last.tail = _ensure_space_between(prev_tail, next_text, after_element=True)


def _append_mixed_content(src: etree._Element, dst: etree._Element) -> None:
    last = dst[-1] if len(dst) > 0 else None
    src_text = src.text or ""
    if last is None:
        existing = dst.text or ""
        dst.text = _ensure_space_between(existing, src_text, after_element=False) + src_text
    else:
        existing = last.tail or ""
        last.tail = _ensure_space_between(existing, src_text, after_element=True) + src_text
    for i, child in enumerate(src):
        copied = copy.deepcopy(child)
        dst.append(copied)
        copied.tail = child.tail or ""
        # Ensure whitespace before next sibling's text unless same word (e.g. un + real)
        next_sib = src[i + 1] if i + 1 < len(src) else None
        next_text = (next_sib.text or "") if next_sib is not None else ""
        if next_text.strip() and not (next_text[0].isspace() if next_text else True):
            if not copied.tail.endswith((" ", "\n", "\t")):
                if not (copied.tail and copied.tail[-1].isalpha() and next_text[0].isalpha()):
                    copied.tail = copied.tail + " "


def _process_paragraph(
    para: Any,
    image_map: dict[str, str],
    hyperlink_map: dict[str, str],
    footnote_map: dict[str, str],
    image_reldir: str,
    run_map: dict,
    tag: str = "p",
) -> Optional[etree._Element]:
    p_elem = etree.Element(tag)
    p_elem.tail = "\n"
    para_style = _paragraph_to_css(para)
    if para_style:
        p_elem.set("style", para_style)

    lasthi = p_elem
    for child in para._element:
        tag = _get_tag(child)
        if tag in ("pPr", "proofErr", "bookmarkStart", "bookmarkEnd", "smartTag"):
            pass
        elif tag == "hyperlink":
            rid = child.get(_qn("r", "id"))
            if rid in hyperlink_map:
                ref_elem = etree.Element("ref")
                ref_elem.set("target", hyperlink_map[rid])
                ref_elem.text = "".join(
                    (n.text or "") for n in child.findall(".//w:t", namespaces=NAMESPACES)
                )
                _ensure_trailing_space_before(p_elem, ref_elem)
                p_elem.append(ref_elem)
                lasthi = ref_elem
        elif tag == "r":
            run = run_map.get(child)
            if run is None:
                continue
            hi = _process_run(run, footnote_map)
            if len(hi) > 0 or (hi.text is not None and hi.text):
                if not hi.attrib:
                    _append_mixed_content(hi, p_elem)
                    lasthi = p_elem
                elif hi.get("style") == lasthi.get("style"):
                    _append_mixed_content(hi, lasthi)
                else:
                    _ensure_trailing_space_before(p_elem, hi)
                    p_elem.append(hi)
                    lasthi = hi
        for drawing in child.findall(".//w:drawing", namespaces=NAMESPACES):
            blip = drawing.find(".//a:blip", namespaces=NAMESPACES)
            if blip is not None:
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed:
                    img_fn = image_map.get(embed)
                    if img_fn:
                        figure = etree.SubElement(p_elem, "figure")
                        figure.set("id", embed)
                        etree.SubElement(figure, "graphic", url=os.path.join(image_reldir, img_fn))

    if len(p_elem) == 0 and not (p_elem.text or "").strip():
        return None
    return p_elem


def _process_table(
    table: Any,
    image_map: dict[str, str],
    hyperlink_map: dict[str, str],
    footnote_map: dict[str, str],
    image_reldir: str,
) -> etree._Element:
    tbl = etree.Element("table")
    for row in table.rows:
        row_el = etree.Element("row")
        for cell in row.cells:
            cell_el = etree.Element("cell")
            for para in cell.paragraphs:
                run_map = {r._element: r for r in para.runs}
                proc = _process_paragraph(
                    para, image_map, hyperlink_map, footnote_map, image_reldir, run_map
                )
                if proc is not None:
                    cell_el.append(proc)
            row_el.append(cell_el)
        tbl.append(row_el)
    return tbl


def _extract_hyperlinks(doc: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    for rId, rel in doc.part.rels.items():
        if hasattr(rel, "target_ref"):
            out[rId] = rel.target_ref
    return out


def _extract_footnotes(doc: Any) -> tuple[Optional[etree._Element], dict[str, str]]:
    footnote_map: dict[str, str] = {}
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in getattr(rel, "target_ref", ""):
            footnotes_part = rel.target_part
            break
    if not footnotes_part:
        return None, footnote_map
    fn_xml = etree.parse(BytesIO(footnotes_part.blob))
    ns = {"w": NAMESPACES["w"]}
    for fn in fn_xml.findall("w:footnote", ns):
        fn_id = fn.get(_qn("w", "id"))
        text = "".join((n.text or "") for n in fn.findall(".//w:t", ns))
        footnote_map[fn_id or ""] = text
    notes_elem = etree.Element("notes")
    for fn_id, fntext in footnote_map.items():
        n = etree.SubElement(notes_elem, "note")
        n.set("id", "fn-" + fn_id)
        n.text = fntext
    if len(notes_elem) == 0 and not (notes_elem.text or "").strip():
        return None, footnote_map
    return notes_elem, footnote_map


def docx_to_tei_tree(
    docx_path: str,
    *,
    orgfile: Optional[str] = None,
    image_dir: Optional[str] = None,
    image_reldir: Optional[str] = None,
) -> tuple[etree._Element, str]:
    """
    Convert a DOCX file to a TEITOK-style TEI element tree (no namespace).
    Returns (tei_root_element, image_dir_used).
    """
    DocxDocument, PackageNotFoundError = _require_docx()
    doc = DocxDocument(docx_path)

    basename = os.path.splitext(os.path.basename(docx_path))[0]
    if image_dir is None:
        image_dir = tempfile.mkdtemp(prefix="flexiconv_docx_")
    if image_reldir is None:
        image_reldir = basename

    image_map = _extract_images_and_map(docx_path, image_dir)
    footnotes_elem, footnote_map = _extract_footnotes(doc)
    hyperlink_map = _extract_hyperlinks(doc)

    if orgfile is None:
        orgfile = docx_path

    XML_NS = "http://www.w3.org/XML/1998/namespace"
    tei = etree.Element("TEI")
    tei_header = etree.SubElement(tei, "teiHeader")
    text_el = etree.SubElement(tei, "text")
    text_el.set(f"{{{XML_NS}}}space", "preserve")
    text_el.set("id", basename)
    body = etree.SubElement(text_el, "body")

    filedesc = etree.SubElement(tei_header, "fileDesc")
    notesstmt = etree.SubElement(filedesc, "notesStmt")
    note = etree.SubElement(notesstmt, "note")
    note.set("n", "orgfile")
    note.text = orgfile
    revisiondesc = etree.SubElement(tei_header, "revisionDesc")
    change = etree.SubElement(revisiondesc, "change")
    change.set("who", "flexiconv")
    change.set("when", str(date.today()))
    change.text = "Converted from DOCX file " + docx_path
    titlestmt = etree.SubElement(filedesc, "titleStmt")
    profiledesc = etree.SubElement(tei_header, "profileDesc")

    meta = doc.core_properties
    if meta.title:
        t = etree.SubElement(titlestmt, "title")
        t.text = meta.title
    if meta.author:
        a = etree.SubElement(titlestmt, "author")
        a.text = meta.author
    if meta.created:
        d = etree.SubElement(titlestmt, "date")
        d.text = str(meta.created)
    if meta.keywords:
        textclass = etree.SubElement(profiledesc, "textClass")
        kw = etree.SubElement(textclass, "keywords")
        term = etree.SubElement(kw, "term")
        term.text = meta.keywords
    if meta.language:
        langusage = etree.SubElement(profiledesc, "langUsage")
        lang = etree.SubElement(langusage, "language")
        lang.text = meta.language

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    def _is_list_paragraph(para: Any) -> bool:
        """Return True if this paragraph is part of a list (bulleted/numbered)."""
        # 1) Check paragraph properties for numbering (OOXML w:pPr/w:numPr)
        pPr = para._element.find(_qn("w", "pPr"))
        if pPr is not None and pPr.find(_qn("w", "numPr")) is not None:
            return True
        # 2) Fallback: python-docx object model if available
        pPr_obj = getattr(para._p, "pPr", None)
        if pPr_obj is not None and getattr(pPr_obj, "numPr", None) is not None:
            return True
        # 3) Style name hints
        name = (getattr(para.style, "name", "") or "").lower()
        return "list" in name or "bullet" in name or "number" in name

    current_list: Optional[etree._Element] = None

    for elem in doc._element.findall(".//w:body/*", namespaces=NAMESPACES):
        tag = _get_tag(elem)
        processed = None
        if tag == "p":
            para = para_map.get(elem)
            if para is not None:
                is_list = _is_list_paragraph(para)
                run_map = {r._element: r for r in para.runs}
                # First paragraph in the document that is not a list and uses a
                # title/heading style becomes a <head>.
                style_name = (getattr(para.style, "name", "") or "").lower()
                want_head = (
                    not is_list
                    and len(body) == 0
                    and ("title" in style_name or style_name.startswith("heading"))
                )
                processed = _process_paragraph(
                    para,
                    image_map,
                    hyperlink_map,
                    footnote_map,
                    image_reldir,
                    run_map,
                    tag="head" if want_head else "p",
                )
                if processed is not None:
                    if is_list:
                        if current_list is None:
                            current_list = etree.SubElement(body, "list")
                        item = etree.SubElement(current_list, "item")
                        # Move children of processed <p> directly under <item>
                        # so <item> contains the text/hi structure without an extra <p>.
                        for child in list(processed):
                            processed.remove(child)
                            item.append(child)
                        if processed.text:
                            if len(item) == 0:
                                item.text = processed.text
                            else:
                                existing = item[-1].tail or ""
                                item[-1].tail = _ensure_space_between(existing, processed.text, after_element=True) + processed.text
                    else:
                        current_list = None
                        body.append(processed)
        elif tag == "tbl":
            current_list = None
            table = table_map.get(elem)
            if table is not None:
                processed = _process_table(
                    table, image_map, hyperlink_map, footnote_map, image_reldir
                )
                if processed is not None:
                    body.append(processed)

    return tei, image_dir


def load_docx(
    path: str,
    *,
    doc_id: Optional[str] = None,
    orgfile: Optional[str] = None,
    image_dir: Optional[str] = None,
    image_reldir: Optional[str] = None,
) -> Document:
    """
    Load a DOCX file into a pivot Document.

    The resulting TEI tree is stored so that saving to TEITOK reproduces that output (header,
    paragraphs, <hi> styles, tables, images, links, footnotes).
    """
    tei_root, used_image_dir = docx_to_tei_tree(
        path,
        orgfile=orgfile or path,
        image_dir=image_dir,
        image_reldir=image_reldir,
    )
    if doc_id is None:
        doc_id = path
    doc = Document(id=doc_id)
    doc.meta["source_filename"] = os.path.basename(path)
    doc.meta["_teitok_tei_root"] = tei_root
    doc.meta["_teitok_image_dir"] = used_image_dir
    return doc


def save_docx(document: Document, path: str, *, source_path: Optional[str] = None, **kwargs: Any) -> None:
    """
    Write a simple DOCX document from the pivot / TEITOK-style TEI.

    This is intentionally conservative:
    - Uses <text>/<body> structure from TEI when available.
    - Maps <head> to Word headings, <p> to paragraphs, <list>/<item> to
      bullet paragraphs, and <table>/<row>/<cell> to Word tables.
    - Inline markup (<hi>, <m>, etc.) is currently flattened to plain text.
    """
    DocxDocument, _ = _require_docx()

    # Obtain a TEI root to drive the structure: prefer stored TEI, otherwise
    # parse the original TEI/XML source if available.
    tei_root = document.meta.get("_teitok_tei_root")
    if tei_root is None and source_path and os.path.isfile(source_path):
        tree = etree.parse(source_path)
        tei_root = tree.getroot()

    if tei_root is None:
        # As a last resort, write a one-line DOCX with a generic message.
        docx = DocxDocument()
        docx.add_paragraph("Converted document (no TEI structure available).")
        docx.save(path)
        return

    # Locate <text>/<body> in the TEI tree.
    text_el = tei_root.xpath(".//*[local-name()='text']")
    text_el = text_el[0] if text_el else tei_root
    body_el = text_el.xpath("./*[local-name()='body']")
    body_el = body_el[0] if body_el else text_el

    def _append_inline_runs(src: etree._Element, para) -> None:
        """
        Append runs to a python-docx paragraph from a TEI element, preserving basic
        inline styling from <hi style='...'> / <hi rend='...'> where possible.
        """

        def _add_text(text: str, bold: Optional[bool], italic: Optional[bool], underline: Optional[bool]) -> None:
            if not text:
                return
            run = para.add_run(text)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline

        def _style_from_hi(el: etree._Element, bold: Optional[bool], italic: Optional[bool], underline: Optional[bool]) -> tuple[Optional[bool], Optional[bool], Optional[bool]]:
            style = (el.get("style") or el.get("rend") or "").lower()
            b, i, u = bold, italic, underline
            if "bold" in style or "font-weight" in style:
                b = True
            if "italic" in style or "emph" in style or "oblique" in style:
                i = True
            if "underline" in style:
                u = True
            # Very simple heuristics for rend='b', 'i', etc.
            if not style:
                rend = (el.get("rend") or "").lower()
                if "b" in rend and b is None:
                    b = True
                if "i" in rend and i is None:
                    i = True
            return b, i, u

        def _walk(node: etree._Element, bold: Optional[bool], italic: Optional[bool], underline: Optional[bool]) -> None:
            if node.text:
                _add_text(node.text, bold, italic, underline)
            for child in node:
                tag = (child.tag or "").split("}")[-1]
                cbold, citalic, cunderline = bold, italic, underline
                if tag == "hi":
                    cbold, citalic, cunderline = _style_from_hi(child, cbold, citalic, cunderline)
                _walk(child, cbold, citalic, cunderline)
                if child.tail:
                    _add_text(child.tail, bold, italic, underline)

        _walk(src, None, None, None)

    docx = DocxDocument()

    # Clear the default empty paragraph if present.
    if len(docx.paragraphs) == 1 and not (docx.paragraphs[0].text or "").strip():
        p = docx.paragraphs[0]
        p.text = ""

    # Walk top-level children under <body>.
    for child in body_el:
        tag = (child.tag or "").split("}")[-1]
        if tag == "head":
            level = _infer_heading_level(child)
            para = docx.add_heading("", level=level)
            _append_inline_runs(child, para)
        elif tag == "list":
            # Each <item> becomes a bullet paragraph.
            for item in child.xpath("./*[local-name()='item']"):
                # Create a bullet paragraph and fill it from the TEI content.
                para = docx.add_paragraph("")
                try:
                    para.style = "List Bullet"
                except Exception:
                    # Fallback if the style does not exist in the template.
                    pass
                _append_inline_runs(item, para)
        elif tag == "table":
            rows = child.xpath("./*[local-name()='row']")
            if not rows:
                continue
            max_cols = 0
            row_cells: list[list[etree._Element]] = []
            for r in rows:
                cells = r.xpath("./*[local-name()='cell']")
                row_cells.append(cells)
                if len(cells) > max_cols:
                    max_cols = len(cells)
            if max_cols == 0:
                continue
            table = docx.add_table(rows=len(row_cells), cols=max_cols)
            for r_idx, cells in enumerate(row_cells):
                for c_idx, cell in enumerate(cells):
                    paragraph = table.rows[r_idx].cells[c_idx].paragraphs[0]
                    # Clear any default text and append runs based on TEI cell content.
                    paragraph.text = ""
                    _append_inline_runs(cell, paragraph)
        else:
            # Default: treat as a paragraph-like block.
            para = docx.add_paragraph("")
            _append_inline_runs(child, para)

    docx.save(path)
